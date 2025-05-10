import tkinter as tk
from tkinter import ttk, filedialog
from docx import Document
from fpdf import FPDF
from styles import SmartStockStyles
from datetime import datetime
import requests
import pytesseract
from tkinter import simpledialog
from tkinter import filedialog
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from tkinter import filedialog
import json
from datetime import datetime
import re
import os
from PIL import Image


class ProductsView:
    def __init__(self, parent, products, save_to_json_callback, show_message_callback):
        self.parent = parent
        self.products = products
        self.save_to_json = save_to_json_callback
        self.show_message = show_message_callback
        self.create_products_ui()
    
    def create_products_ui(self):
        """Create the products UI"""
        self.create_header()
        
        content_container = ttk.Frame(self.parent, style="Main.TFrame")
        content_container.pack(fill="both", expand=True, padx=30, pady=(0, 30))
        
        form_frame = ttk.Frame(content_container, style="Main.TFrame")
        form_frame.pack(side="left", fill="both", expand=True, padx=(0, 15))
        self.create_product_form(form_frame)
        
        products_frame = ttk.Frame(content_container, style="Main.TFrame")
        products_frame.pack(side="right", fill="both", expand=True, padx=(15, 0))
        self.create_products_display(products_frame)
    
    def create_header(self):
        """Create the header section"""
        header_frame = ttk.Frame(self.parent, style="Main.TFrame")
        header_frame.pack(fill="x", padx=30, pady=(30, 20))

        
        title_label = ttk.Label(
            header_frame,
            text="Product Manager",
            style="Title.TLabel"
        )
        title_label.pack(side="left")
        
        subtitle_label = ttk.Label(
            header_frame,
            text="Manage your inventory with ease",
            font=("Segoe UI", 12),
            background=SmartStockStyles.COLORS["bg_light"],
            foreground=SmartStockStyles.COLORS["text_muted"]
        )
        subtitle_label.pack(side="left", padx=(10, 0), pady=5)
        
        date_label = ttk.Label(
            header_frame,
            text=datetime.now().strftime("%B %d, %Y"),
            font=("Segoe UI", 12),
            background=SmartStockStyles.COLORS["bg_light"],
            foreground=SmartStockStyles.COLORS["text_muted"]
        )
        date_label.pack(side="right", pady=5)
    
    def create_product_form(self, parent):
        """Create the product input form"""
        section_title = ttk.Label(
            parent,
            text="Add New Product",
            style="Subtitle.TLabel"
        )
        section_title.pack(anchor="w", pady=(0, 15))
        
        form_card = ttk.Frame(
            parent,
            style="Card.TFrame",
            padding=SmartStockStyles.SPACING["lg"]
        )
        form_card.pack(fill="both", expand=True)
        
        name_label = ttk.Label(
            form_card,
            text="Product Name",
            font=SmartStockStyles.FONTS["body"],
            background=SmartStockStyles.COLORS["bg_card"],
            foreground=SmartStockStyles.COLORS["text_dark"]
        )
        name_label.pack(anchor="w", pady=(0, 5))
        
        self.entry_name = tk.Entry(form_card)
        SmartStockStyles.apply_entry_style(self.entry_name)
        self.entry_name.pack(fill="x", pady=(0, 15))
        
        quantity_label = ttk.Label(
            form_card,
            text="Quantity",
            font=SmartStockStyles.FONTS["body"],
            background=SmartStockStyles.COLORS["bg_card"],
            foreground=SmartStockStyles.COLORS["text_dark"]
        )
        quantity_label.pack(anchor="w", pady=(0, 5))
        
        self.entry_quantity = tk.Entry(form_card)
        SmartStockStyles.apply_entry_style(self.entry_quantity)
        self.entry_quantity.pack(fill="x", pady=(0, 15))
        
        price_label = ttk.Label(
            form_card,
            text="Price ($)",
            font=SmartStockStyles.FONTS["body"],
            background=SmartStockStyles.COLORS["bg_card"],
            foreground=SmartStockStyles.COLORS["text_dark"]
        )
        price_label.pack(anchor="w", pady=(0, 5))
        
        self.entry_price = tk.Entry(form_card)
        SmartStockStyles.apply_entry_style(self.entry_price)
        self.entry_price.pack(fill="x", pady=(0, 15))
        
        add_button = tk.Button(
            form_card,
            text="Add Product",
            command=self.add_product
        )
        SmartStockStyles.apply_button_style(add_button, "primary")
        add_button.pack(fill="x", pady=(15, 0), ipady=8)
        
        button_frame = ttk.Frame(form_card, style="Card.TFrame")
        button_frame.pack(fill="x", pady=(20, 0))
        
        button_frame.columnconfigure(0, weight=1)
        button_frame.columnconfigure(1, weight=1)

        
        # --- Export to PDF Button ---
        pdf_button = tk.Button(
            button_frame,
            text="Export to PDF",
            command=self.export_to_pdf
        )
        SmartStockStyles.apply_button_style(pdf_button, "outline")
        pdf_button.grid(row=0, column=0, padx=(0, 5), pady=(0, 5), sticky="ew")

        # --- Generate AI Report Button ---
        ai_report = tk.Button(
            button_frame,
            text="📃 Generate AI Report",
            command=self.generate_ai_report,
            bg="#4CAF50",
            fg="white",
            font=("Segoe UI", 10),
        )
        SmartStockStyles.apply_button_style(ai_report, "secondary")
        ai_report.grid(row=0, column=1, padx=(5, 0), pady=(0, 5), sticky="ew")

        # --- Save to Word Button ---
        word_button = tk.Button(
            button_frame,
            text="Save to Word",
            command=self.save_to_word
        )
        SmartStockStyles.apply_button_style(word_button, "outline")
        word_button.grid(row=1, column=0, padx=(0, 5), pady=(5, 0), sticky="ew")

        # --- AI Assistant Button ---
        ai_button = tk.Button(
            button_frame,
            text="AI Assistant 🤖",
            command=self.show_ai_assistant
        )
        SmartStockStyles.apply_button_style(ai_button, "secondary")
        ai_button.grid(row=1, column=1, padx=(5, 0), pady=(5, 0), sticky="ew")

        # --- Load from Word Button ---
        load_button = tk.Button(
            button_frame,
            text="Load from Word",
            command=self.load_from_word
        )
        SmartStockStyles.apply_button_style(load_button, "outline")
        load_button.grid(row=2, column=0, padx=(0, 5), pady=(5, 0), sticky="ew")

        # --- Tracking Daily Sales ---
        tracking_button = tk.Button(
            button_frame,
            text="📊 Today Tracking",
            command=self.summarize_today_activity
        )
        SmartStockStyles.apply_button_style(tracking_button, "secondary")
        tracking_button.grid(row=2, column=1, padx=(5, 0), pady=(5, 0), sticky="ew")


    def create_products_display(self, parent):
        """Create the products display section"""
        header_frame = ttk.Frame(parent, style="Main.TFrame")
        header_frame.pack(fill="x", pady=(0, 15))
        
        section_title = ttk.Label(
            header_frame,
            text="Product List",
            style="Subtitle.TLabel"
        )
        section_title.pack(side="left")
        
        search_frame = ttk.Frame(header_frame, style="Main.TFrame")
        search_frame.pack(side="right")
        
        self.search_var = tk.StringVar()
        search_entry = tk.Entry(
            search_frame,
            textvariable=self.search_var,
            width=20
        )
        SmartStockStyles.apply_entry_style(search_entry)
        search_entry.pack(side="left", padx=(0, 5))
        
        search_button = tk.Button(
            search_frame,
            text="🔍 Search",
            command=self.search_product
        )
        SmartStockStyles.apply_button_style(search_button, "secondary")
        search_button.pack(side="right")
        
        table_card = ttk.Frame(
            parent,
            style="Card.TFrame",
            padding=(10, 10, 10, 10)
        )
        table_card.pack(fill="both", expand=True)
        
        self.tree = ttk.Treeview(
            table_card, 
            columns=("Name", "Quantity", "Price"), 
            show="headings",
            height=12
        )
        
        self.tree.heading("Name", text="Product Name")
        self.tree.heading("Quantity", text="Quantity")
        self.tree.heading("Price", text="Price")
        
        self.tree.column("Name", width=150, anchor="w")
        self.tree.column("Quantity", width=80, anchor="center")
        self.tree.column("Price", width=80, anchor="e")
        
        scrollbar = ttk.Scrollbar(table_card, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        self.tree.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Add delete button
        delete_button = tk.Button(
            table_card,
            text="🗑️ Delete Selected",
            command=self.delete_selected_product
        )
        SmartStockStyles.apply_button_style(delete_button, "error")
        delete_button.pack(pady=10)
        
        # Add update button
        update_button = tk.Button(
            table_card,
            text="✏️ Update Selected",
            command=self.update_selected_product
        )
        SmartStockStyles.apply_button_style(update_button, "secondary")
        update_button.pack(pady=(0, 10))
        
        status_frame = ttk.Frame(parent, style="Main.TFrame")
        status_frame.pack(fill="x", pady=(10, 0))
        
        self.status_label = ttk.Label(
            status_frame,
            text="0 products in inventory",
            font=SmartStockStyles.FONTS["small"],
            background=SmartStockStyles.COLORS["bg_light"],
            foreground=SmartStockStyles.COLORS["text_muted"]
        )
        self.status_label.pack(side="left")
        
        self.update_tree()
    
    def add_product(self):
        """Add a product to the inventory"""
        name = self.entry_name.get().strip()
        quantity = self.entry_quantity.get().strip()
        price = self.entry_price.get().strip()
        
        if not name:
            self.show_message("Please enter a product name", "error")
            return
            
        try:
            quantity = int(quantity)
            if quantity < 0:
                self.show_message("Quantity must be a positive number", "error")
                return
        except ValueError:
            self.show_message("Quantity must be a valid number", "error")
            return
            
        try:
            price = float(price)
            if price < 0:
                self.show_message("Price must be a positive number", "error")
                return
        except ValueError:
            self.show_message("Price must be a valid number", "error")
            return
        
        product = {"name": name, "quantity": quantity, "price": price}
        self.products.append(product)
        
        self.update_tree()
        self.save_to_json()
        
        self.entry_name.delete(0, tk.END)
        self.entry_quantity.delete(0, tk.END)
        self.entry_price.delete(0, tk.END)
        
        self.show_message(f"Product '{name}' added successfully", "success")
    
    def update_tree(self):
        """Update the treeview with current products"""
        self.tree.delete(*self.tree.get_children())
        
        for p in self.products:
            self.tree.insert(
                "", 
                "end", 
                values=(p["name"], p["quantity"], f"${p['price']:.2f}")
            )
        
        count = len(self.products)
        self.status_label.config(
            text=f"{count} {'product' if count == 1 else 'products'} in inventory"
        )

    def save_to_word(self):
        """Save products to a Word document"""
        if not self.products:
            self.show_message("No products to save", "warning")
            return
            
        try:
            doc = Document()
            doc.add_heading("SmartStock - Product List", 0)
            
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
            doc.add_paragraph("")
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Product Name'
            header_cells[1].text = 'Quantity'
            header_cells[2].text = 'Price'
            
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            for p in self.products:
                row_cells = table.add_row().cells
                row_cells[0].text = p["name"]
                row_cells[1].text = str(p["quantity"])
                row_cells[2].text = f"${p['price']:.2f}"
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Save Product List As",
                initialfile="products.docx"
            )
            
            if file_path:
                doc.save(file_path)
                self.show_message(f"Products saved to {os.path.basename(file_path)}", "success")
            
        except Exception as e:
            self.show_message(f"Error saving to Word: {str(e)}", "error")
    
    def export_to_pdf(self):
        """Export products to a PDF document"""
        if not self.products:
            self.show_message("No products to export", "warning")
            return
            
        try:
            pdf = FPDF()
            pdf.add_page()
            
            pdf.set_font("Arial", size=12)
            
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(200, 10, txt="SmartStock - Product List", ln=True, align='C')
            
            pdf.set_font("Arial", '', 10)
            pdf.cell(200, 10, txt=f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}", ln=True)
            pdf.ln(5)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(90, 10, "Product Name", border=1)
            pdf.cell(40, 10, "Quantity", border=1, align='C')
            pdf.cell(60, 10, "Price", border=1, align='R')
            pdf.ln()
            
            pdf.set_font("Arial", '', 12)
            for p in self.products:
                pdf.cell(90, 10, p["name"], border=1)
                pdf.cell(40, 10, str(p["quantity"]), border=1, align='C')
                pdf.cell(60, 10, f"${p['price']:.2f}", border=1, align='R')
                pdf.ln()
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
                title="Save Product List As",
                initialfile="products.pdf"
            )
            
            if file_path:
                pdf.output(file_path)
                self.show_message(f"Products exported to {os.path.basename(file_path)}", "success")
            
        except Exception as e:
            self.show_message(f"Error exporting to PDF: {str(e)}", "error")
    
    def load_from_word(self):
        """Load products from a Word document"""
        try:
            file_path = filedialog.askopenfilename(
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Open Product List",
                initialfile="products.docx"
            )
            
            if not file_path:
                return
                
            doc = Document(file_path)
            tables = doc.tables
            
            if not tables:
                self.show_message("No table found in the Word document", "warning")
                return
                
            table = tables[0]
            self.products.clear()
            
            for row in table.rows[1:]:
                try:
                    name = row.cells[0].text.strip()
                    quantity = int(row.cells[1].text.strip())
                    price = float(row.cells[2].text.strip().replace('$', ''))
                    
                    self.products.append({
                        "name": name,
                        "quantity": quantity,
                        "price": price
                    })
                except (ValueError, IndexError) as e:
                    self.show_message(f"Skipping invalid row: {str(e)}", "warning")
                    continue
            
            self.update_tree()
            self.save_to_json()
            self.show_message(f"Products loaded from {os.path.basename(file_path)}", "success")
            
        except Exception as e:
            self.show_message(f"Error loading products: {str(e)}", "error")
    
    def search_product(self):
        """Search for a product by name"""
        search_term = self.search_var.get().strip().lower()
        
        self.tree.delete(*self.tree.get_children())
        
        if not search_term:
            for p in self.products:
                self.tree.insert(
                    "", 
                    "end", 
                    values=(p["name"], p["quantity"], f"${p['price']:.2f}")
                )
            count = len(self.products)
            self.status_label.config(
                text=f"{count} {'product' if count == 1 else 'products'} in inventory"
            )
            return
            
        found = False
        for p in self.products:
            if search_term in p["name"].lower():
                self.tree.insert(
                    "", 
                    "end", 
                    values=(p["name"], p["quantity"], f"${p['price']:.2f}")
                )
                found = True
        
        count = len(self.tree.get_children())
        self.status_label.config(
            text=f"{count} {'result' if count == 1 else 'results'} found"
        )
        if not found:
            self.show_message("No products matching your search", "warning")


    def save_report_to_pdf(self, report_text):
        """Save the generated report to a PDF file"""
        try:
            # Open a file dialog to choose where to save the PDF
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF Files", "*.pdf")],
                title="Save Report as PDF"
            )

            if file_path:
                # Create a PDF document
                c = canvas.Canvas(file_path, pagesize=letter)

                # Add the report text to the PDF
                text_object = c.beginText(40, 750) 
                text_object.setFont("Helvetica", 10)
                text_object.setTextOrigin(40, 750)

                # Split the report into lines
                lines = report_text.split("\n")
                for line in lines:
                    text_object.textLine(line)

                # Draw the text on the PDF
                c.drawText(text_object)

                # Save the PDF
                c.showPage()
                c.save()

                # Inform the user that the PDF has been saved
                self.show_message("Report saved as PDF successfully!", "success")

        except Exception as e:
            self.show_message(f"Error saving PDF: {str(e)}", "error")


    def summarize_today_activity(self):
        """Use SmartStock AI to summarize today's stock activity from"""
        try:
            # --- Load and filter today's invoices ---
            with open("invoices.json", "r", encoding="utf-8") as f:
                invoices = json.load(f)

            today = datetime.today().date()
            sales_today = []

            for invoice in invoices:
                date_str = invoice.get("date")
                if not date_str:
                    continue

                try:
                    invoice_date = datetime.strptime(date_str, "%Y-%m-%d").date()
                except ValueError:
                    continue

                if invoice_date == today:
                    sales_today.append(invoice)

            if not sales_today:
                summary_text = "No sales were made today."
            else:
                # --- Build sales summary text ---
                lines = []
                for sale in sales_today:
                    client = sale.get("client_name", "Unknown Client")
                    for product in sale.get("products", []):
                        name = product.get("name", "Unknown Product")
                        qty = product.get("quantity", 1)
                        lines.append(f"- Sold {qty} × {name} to {client}")

                sales_description = "\n".join(lines)

                # --- Ask AI for summary ---
                prompt = (
                    f"Based on the following sales made today:\n"
                    f"{sales_description}\n\n"
                    f"Write a short summary (2-3 sentences) of today's stock activity: total items sold, client names, and highlight any product sold the most."
                )

                response = requests.post(
                    "http://localhost:11434/api/generate",
                    json={
                        "model": "llama3",
                        "prompt": prompt,
                        "stream": False
                    }
                )

                summary_text = response.json()["response"].strip() if response.ok else "AI unavailable"

            # --- Show result in popup ---
            popup = tk.Toplevel(self.parent)
            popup.title("📊 Today’s Sales Summary")
            popup.configure(bg=SmartStockStyles.COLORS["bg_card"], padx=30, pady=30)
            popup.resizable(True, True)

            text_widget = tk.Text(popup, wrap="word", font=("Segoe UI", 10), bg="#f8f8f8", height=12, width=70)
            text_widget.insert("1.0", summary_text)
            text_widget.config(state="disabled")
            text_widget.pack()

            tk.Button(popup, text="✕ Close", command=popup.destroy).pack(pady=10)
            popup.transient(self.parent)
            popup.grab_set()

        except Exception as e:
            self.show_message(f"AI Summary Error: {str(e)}", "error")


    def generate_ai_report(self):
        """Generate a natural language inventory report using Ollama AI"""
        try:
            # Show loading message with RPG styling
            self.show_message("The ancient scroll is being prepared...", "info")

            # --- Build product summary for context ---
            product_summary = "\n".join([
                f"- {p['name']}: {p['quantity']} in stock"
                for p in self.products
            ])

            # --- Prompt for AI ---
            prompt = (
                f"Based on the following product inventory:\n{product_summary}\n\n"
                f"Write a professional monthly inventory report including:\n"
                f"1. Summary of stock levels\n"
                f"2. Products that are low in stock\n"
                f"3. Any insights you can give\n"
                f"4. the name of inventory manager is mouhssine\n"
                f"Write in clear and formal language."
            )

            # --- Send to Ollama ---
            response = requests.post(
                "http://localhost:11434/api/generate",
                json={
                    "model": "llama3",
                    "prompt": prompt,
                    "stream": False
                }
            )

            report_text = response.json()["response"].strip() if response.ok else "The ancient scroll appears to be blank. (AI unavailable)"

            # --- Show in popup ---
            popup = tk.Toplevel(self.parent)
            popup.title("📜 Ancient Inventory Scroll")
            popup.resizable(True, True)

            # RPG dark background color
            popup.configure(bg="#1a1c42", padx=30, pady=30)

            # Center the popup window
            window_width = 900
            window_height = 900
            screen_width = popup.winfo_screenwidth()
            screen_height = popup.winfo_screenheight()
            center_x = int(screen_width/2 - window_width/2)
            center_y = int(screen_height/2 - window_height/2)
            popup.geometry(f'{window_width}x{window_height}+{center_x}+{center_y}')

            # Create main scroll frame
            scroll_frame = tk.Frame(
                popup,
                bg="#f0e6d2",  # Parchment color
                highlightbackground="#c8aa6e",  # Gold border
                highlightthickness=3,
                bd=0
            )
            scroll_frame.pack(fill="both", expand=True, padx=10, pady=10)

            # Header frame
            header_frame = tk.Frame(scroll_frame, bg="#3c275a", height=40)
            header_frame.pack(fill="x")

            # Title with RPG styling
            title_label = tk.Label(
                header_frame,
                text="✧ Inventory Chronicles ✧",
                font=("Palatino Linotype", 16, "bold"),
                bg="#3c275a",  # Royal purple
                fg="#e5c770"   # Golden text
            )
            title_label.pack(pady=10)

            # Decorative separator
            separator_frame = tk.Frame(scroll_frame, height=10, bg="#f0e6d2")
            separator_frame.pack(fill="x")

            # Create a canvas with scrollbar for the report
            canvas_frame = tk.Frame(scroll_frame, bg="#f0e6d2")
            canvas_frame.pack(fill="both", expand=True, padx=15, pady=5)

            scrollbar = tk.Scrollbar(canvas_frame)
            scrollbar.pack(side="right", fill="y")

            # Report display with parchment styling
            report_label = tk.Text(
                canvas_frame, 
                wrap="word", 
                font=("Palatino Linotype", 11), 
                bg="#f5f2e9",  # Lighter parchment for text area
                fg="#483248",  # Dark purple text
                padx=20, 
                pady=20, 
                height=20, 
                width=70,
                relief="flat",
                bd=0,
                highlightthickness=0,
                yscrollcommand=scrollbar.set
            )
            report_label.insert("1.0", report_text)
            report_label.config(state="disabled")
            report_label.pack(side="left", fill="both", expand=True)
            scrollbar.config(command=report_label.yview)

            # Add a footer with date
            footer_frame = tk.Frame(scroll_frame, bg="#f0e6d2", height=30)
            footer_frame.pack(fill="x", pady=(0, 5))

            # Simulated wax seal image using a label
            seal_label = tk.Label(
                footer_frame,
                text="🔮",  # Crystal ball emoji as magical seal
                font=("Arial", 16),
                bg="#f0e6d2",
                fg="#8b4513"  # Brown seal color
            )
            seal_label.pack(side="left", padx=20)

            # Date in fantasy style
            import datetime
            current_date = datetime.datetime.now()
            month_names = ["January", "February", "March", "April", "May", "June", 
                          "July", "August", "September", "October", "November", "December"]
            fantasy_date = f"Issued on the {current_date.day}th day of {month_names[current_date.month-1]}, Year {current_date.year}"

            date_label = tk.Label(
                footer_frame,
                text=fantasy_date,
                font=("Palatino Linotype", 9, "italic"),
                bg="#f0e6d2",
                fg="#6b4226"  # Dark brown text
            )
            date_label.pack(side="right", padx=20)

            # Button container frame with parchment background
            button_frame = tk.Frame(scroll_frame, bg="#f0e6d2")
            button_frame.pack(pady=15, padx=15, fill="x", expand=True)

            # Configure grid for button placement
            button_frame.grid_columnconfigure(0, weight=1)
            button_frame.grid_columnconfigure(1, weight=1)

            # Save PDF button (styled as magical scroll)
            save_pdf_button = tk.Button(
                button_frame,
                text="✦ Save To PDF✦",
                command=lambda: self.save_report_to_pdf(report_text),
                bg="#5e3a98",  # Purple
                fg="#f1d9a7",  # Light gold text
                activebackground="#7649b7",  # Lighter purple when pressed
                font=("Palatino Linotype", 10, "bold"),
                padx=20,
                pady=5,
                borderwidth=0,
                relief=tk.RIDGE,
                cursor="hand2"
            )
            save_pdf_button.grid(row=0, column=0, padx=(0, 5), sticky="ew")

            # Close button (styled as dark magic)
            close_button = tk.Button(
                button_frame,
                text="✧ Close Scroll ✧",
                command=popup.destroy,
                bg="#6b4f2d",  # Brown
                fg="#f0e6d2",  # Parchment color
                activebackground="#8c6640",  # Lighter brown when pressed
                font=("Palatino Linotype", 10, "bold"),
                padx=20,
                pady=5,
                borderwidth=0,
                relief=tk.RIDGE,
                cursor="hand2"
            )
            close_button.grid(row=0, column=1, padx=(5, 0), sticky="ew")

            # Add decorative elements - magical runes
            left_rune = tk.Label(
                popup,
                text="⚜",
                font=("Arial", 16),
                bg="#1a1c42",
                fg="#c8aa6e"  # Gold
            )
            left_rune.place(x=15, y=15)

            right_rune = tk.Label(
                popup,
                text="⚜",
                font=("Arial", 16),
                bg="#1a1c42",
                fg="#c8aa6e"  # Gold
            )
            right_rune.place(relx=1.0, x=-30, y=15)

            popup.transient(self.parent)
            popup.grab_set()

        except Exception as e:
            self.show_message(f"The scroll could not be revealed: {str(e)}", "error")


    def show_ai_assistant(self):
            """AI Assistant to answer questions about your product inventory"""
            try:
                # --- Custom RPG-styled Question Dialog ---
                # Create a custom dialog instead of using simpledialog
                question_dialog = tk.Toplevel(self.parent)
                question_dialog.title("Mystic Oracle")
                question_dialog.iconbitmap() if hasattr(question_dialog, 'iconbitmap') else None  # Remove icon on some platforms
                question_dialog.configure(bg="#1a1c42")  # RPG dark blue background
                question_dialog.resizable(False, False)

                # Center the dialog window
                window_width = 450
                window_height = 350
                screen_width = question_dialog.winfo_screenwidth()
                screen_height = question_dialog.winfo_screenheight()
                center_x = int(screen_width/2 - window_width/2)
                center_y = int(screen_height/2 - window_height/2)
                question_dialog.geometry(f'{window_width}x{window_height}+{center_x}+{center_y}')

                # Dialog frame with scroll-like appearance
                dialog_frame = tk.Frame(
                    question_dialog,
                    bg="#f0e6d2",  # Parchment color
                    highlightbackground="#c8aa6e",  # Gold border
                    highlightthickness=2
                )
                dialog_frame.pack(pady=20, padx=20, fill="both", expand=True)

                # Title for the dialog
                title_label = tk.Label(
                    dialog_frame,
                    text="✨ Ask the Oracle ✨",
                    font=("Palatino Linotype", 14, "bold"),
                    bg="#3c275a",  # Royal purple
                    fg="#e5c770",   # Golden text
                    padx=10,
                    pady=5
                )
                title_label.pack(fill="x")

                # Instruction label
                instruction_label = tk.Label(
                    dialog_frame,
                    text="What knowledge do you seek about your inventory?\nE.g. 'What product has the highest quantity?'",
                    font=("Palatino Linotype", 10),
                    bg="#f0e6d2",
                    fg="#483248"
                )
                instruction_label.pack(pady=(15, 5))

                # Text entry for the question
                question_entry = tk.Entry(
                    dialog_frame,
                    font=("Palatino Linotype", 11),
                    width=40,
                    bg="#fffbf0",  # Light parchment
                    fg="#483248",  # Dark purple text
                    insertbackground="#483248",  # Cursor color
                    relief="sunken",
                    bd=1
                )
                question_entry.pack(pady=10, padx=20, ipady=3)
                question_entry.focus_set()  # Set focus to the entry

                # Store the result
                user_question_result = [None]

                # Button functions
                def on_ok():
                    user_question_result[0] = question_entry.get()
                    question_dialog.destroy()

                def on_cancel():
                    question_dialog.destroy()

                # Button frame
                button_frame = tk.Frame(dialog_frame, bg="#f0e6d2")
                button_frame.pack(pady=10)

                # OK button - RPG styled
                ok_button = tk.Button(
                    button_frame,
                    text="Consult",
                    command=on_ok,
                    font=("Palatino Linotype", 10, "bold"),
                    bg="#6b4f2d",  # Brown leather button
                    fg="#f1d9a7",  # Light gold text
                    activebackground="#8c6e4b",
                    activeforeground="#fff9e8",
                    padx=15,
                    pady=3,
                    relief="raised",
                    bd=1
                )
                ok_button.pack(side="left", padx=10)

                # Cancel button - RPG styled
                cancel_button = tk.Button(
                    button_frame,
                    text="Decline",
                    command=on_cancel,
                    font=("Palatino Linotype", 10),
                    bg="#3c275a",  # Royal purple
                    fg="#d9d0c1",  # Light parchment color
                    activebackground="#513969",
                    activeforeground="#f0e6d2",
                    padx=15,
                    pady=3,
                    relief="raised",
                    bd=1
                )
                cancel_button.pack(side="left", padx=10)

                # Add decorative elements
                left_rune = tk.Label(
                    question_dialog,
                    text="⚜",
                    font=("Arial", 16),
                    bg="#1a1c42",
                    fg="#c8aa6e"  # Gold
                )
                left_rune.place(x=15, y=15)

                right_rune = tk.Label(
                    question_dialog,
                    text="⚜",
                    font=("Arial", 16),
                    bg="#1a1c42",
                    fg="#c8aa6e"  # Gold
                )
                right_rune.place(relx=1.0, x=-30, y=15)

                # Make dialog modal
                question_dialog.transient(self.parent)
                question_dialog.grab_set()
                question_dialog.wait_window()

                # Get the result
                user_question = user_question_result[0]

                if not user_question:
                    return  # Cancelled or empty input

                # --- Format Products into Prompt ---
                product_lines = [f"- {p['name']} (Quantity: {p['quantity']})" for p in self.products]
                product_data = "\n".join(product_lines)

                prompt = f"""You are a smart inventory assistant. Here is a list of products:{product_data} Answer the question: {user_question}"""

                # --- Get AI Response (Ollama - llama3) ---
                response = requests.post(
                    "http://localhost:11434/api/generate",
                    json={
                        "model": "llama3",
                        "prompt": prompt,
                        "stream": False
                    }
                )

                ai_answer = response.json()["response"].strip() if response.ok else "AI unavailable."

                # --- Create Popup with AI Answer ---
                popup = tk.Toplevel(self.parent)
                popup.title("Mystic Oracle")
                popup.resizable(False, False)

                # RPG-style background color (deep midnight blue)
                popup.configure(
                    bg="#1a1c42",  # Deep midnight blue background for fantasy theme
                    padx=50,
                    pady=50
                )

                # Center the popup window
                window_width = 450
                window_height = 500
                screen_width = popup.winfo_screenwidth()
                screen_height = popup.winfo_screenheight()
                center_x = int(screen_width/2 - window_width/2)
                center_y = int(screen_height/2 - window_height/2)
                popup.geometry(f'{window_width}x{window_height}+{center_x}+{center_y}')

                # Frame for the answer with parchment-like background
                answer_frame = tk.Frame(
                    popup,
                    bg="#f0e6d2",  # Parchment/scroll color
                    highlightbackground="#c8aa6e",  # Gold border
                    highlightthickness=2,
                    bd=0
                )
                answer_frame.pack(pady=10, padx=5, fill="both", expand=True)

                # Title banner for the answer
                title_frame = tk.Frame(
                    answer_frame,
                    bg="#3c275a",  # Royal purple
                    height=30
                )
                title_frame.pack(fill="x")

                title_label = tk.Label(
                    title_frame,
                    text="✨ Mystic Oracle ✨",  # RPG-themed name instead of "AI Answer"
                    font=("Palatino Linotype", 12, "bold"),
                    bg="#3c275a",  # Royal purple
                    fg="#e5c770"   # Golden text
                )
                title_label.pack(pady=5)

                # Answer text with scroll-like appearance
                msg_label = tk.Text(
                    answer_frame,
                    wrap="word",
                    width=40,
                    height=12,
                    font=("Palatino Linotype", 10),
                    bg="#f0e6d2",  # Parchment color
                    fg="#483248",  # Dark purple text
                    bd=0,
                    padx=15,
                    pady=15,
                    relief="flat"
                )
                msg_label.insert("1.0", ai_answer)
                msg_label.config(state="disabled")  # Make read-only
                msg_label.pack(pady=10, padx=10, fill="both", expand=True)

                # RPG-styled close button
                close_btn = tk.Button(
                    popup,
                    text="✕ Close Scroll",
                    command=popup.destroy,
                    font=("Palatino Linotype", 10, "bold"),
                    bg="#6b4f2d",  # Brown leather button
                    fg="#f1d9a7",  # Light gold text
                    bd=0,
                    relief=tk.FLAT,
                    activebackground="#8c6e4b",  # Lighter brown on hover
                    activeforeground="#fff9e8",  # Brighter text on hover
                    padx=15,
                    pady=5,
                    cursor="hand2"
                )
                close_btn.pack(pady=15)

                # Add decorative elements - magical runes (using simple symbols)
                left_rune = tk.Label(
                    popup,
                    text="⚜",
                    font=("Arial", 16),
                    bg="#1a1c42",
                    fg="#c8aa6e"  # Gold
                )
                left_rune.place(x=15, y=15)

                right_rune = tk.Label(
                    popup,
                    text="⚜",
                    font=("Arial", 16),
                    bg="#1a1c42",
                    fg="#c8aa6e"  # Gold
                )
                right_rune.place(relx=1.0, x=-30, y=15)

                popup.transient(self.parent)
                popup.grab_set()

            except Exception as e:
                self.show_message(f"AI Error: {str(e)}", "error")

            
    def update_products(self, products):
        """Update products and refresh UI"""
        self.products = products
        self.update_tree()
    
    def delete_selected_product(self):
        """Delete the selected product from the inventory and data.json"""
        selected_item = self.tree.selection()
        if not selected_item:
            self.show_message("Please select a product to delete", "warning")
            return
        
        # Get the product name from the selected item
        product_name = self.tree.item(selected_item[0])['values'][0]
        
        # Create a new list excluding the selected product
        new_products = [p for p in self.products if p["name"] != product_name]
        
        # Update the products list only if the save is successful
        try:
            self.products[:] = new_products  # Update the list in-place
            self.save_to_json()  # Save to data.json
            self.update_tree()   # Refresh the Treeview
            self.show_message(f"Product '{product_name}' deleted successfully", "success")
        except Exception as e:
            self.show_message(f"Error deleting product: {str(e)}", "error")
            self.update_tree()  # Refresh Treeview to reflect unchanged products
    
    def update_selected_product(self):
        """Update the selected product in the inventory and data.json"""
        selected_item = self.tree.selection()
        if not selected_item:
            self.show_message("Please select a product to update", "warning")
            return
            
        # Get the current product details
        current_values = self.tree.item(selected_item[0])['values']
        current_name = current_values[0]
        
        # Create update popup
        popup = tk.Toplevel(self.parent)
        popup.title("Update Product")
        popup.geometry("500x600")
        popup.configure(bg=SmartStockStyles.COLORS["bg_light"])
        popup.transient(self.parent)
        popup.minsize(500, 600)
        
        # Center the popup
        screen_width = popup.winfo_screenwidth()
        screen_height = popup.winfo_screenheight()
        x = (screen_width - 500) // 2
        y = (screen_height - 600) // 2
        popup.geometry(f"500x600+{x}+{y}")
        
        # Create form frame with enhanced styling
        form_frame = ttk.Frame(popup, style="Card.TFrame", padding=40)
        form_frame.pack(fill="both", expand=True, padx=40, pady=40)
        
        # Title with enhanced styling
        title_label = ttk.Label(
            form_frame,
            text="Update Product",
            style="Title.TLabel",
            font=("Segoe UI", 24, "bold")
        )
        title_label.pack(anchor="w", pady=(0, 30))
        
        # Name field with enhanced styling
        name_label = ttk.Label(
            form_frame,
            text="Product Name",
            font=("Segoe UI", 12, "bold"),
            background=SmartStockStyles.COLORS["bg_card"],
            foreground=SmartStockStyles.COLORS["text_dark"]
        )
        name_label.pack(anchor="w", pady=(0, 10))
        
        name_entry = tk.Entry(form_frame, font=("Segoe UI", 12))
        SmartStockStyles.apply_entry_style(name_entry)
        name_entry.insert(0, current_values[0])
        name_entry.pack(fill="x", pady=(0, 25))
        
        # Quantity field with enhanced styling
        quantity_label = ttk.Label(
            form_frame,
            text="Quantity",
            font=("Segoe UI", 12, "bold"),
            background=SmartStockStyles.COLORS["bg_card"],
            foreground=SmartStockStyles.COLORS["text_dark"]
        )
        quantity_label.pack(anchor="w", pady=(0, 10))
        
        quantity_entry = tk.Entry(form_frame, font=("Segoe UI", 12))
        SmartStockStyles.apply_entry_style(quantity_entry)
        quantity_entry.insert(0, current_values[1])
        quantity_entry.pack(fill="x", pady=(0, 25))
        
        # Price field with enhanced styling
        price_label = ttk.Label(
            form_frame,
            text="Price ($)",
            font=("Segoe UI", 12, "bold"),
            background=SmartStockStyles.COLORS["bg_card"],
            foreground=SmartStockStyles.COLORS["text_dark"]
        )
        price_label.pack(anchor="w", pady=(0, 10))
        
        price_entry = tk.Entry(form_frame, font=("Segoe UI", 12))
        SmartStockStyles.apply_entry_style(price_entry)
        price_entry.insert(0, current_values[2].replace('$', ''))
        price_entry.pack(fill="x", pady=(0, 30))
        
        def save_update():
            try:
                # Get new values
                new_name = name_entry.get().strip()
                new_quantity = int(quantity_entry.get().strip())
                new_price = float(price_entry.get().strip())
                
                # Validate inputs
                if not new_name:
                    self.show_message("Please enter a product name", "error")
                    return
                    
                if new_quantity < 0:
                    self.show_message("Quantity must be a positive number", "error")
                    return
                    
                if new_price < 0:
                    self.show_message("Price must be a positive number", "error")
                    return
                
                # Update the product
                for product in self.products:
                    if product["name"] == current_name:
                        product["name"] = new_name
                        product["quantity"] = new_quantity
                        product["price"] = new_price
                        break
                
                # Save to JSON and update UI
                self.save_to_json()
                self.update_tree()
                self.show_message(f"Product updated successfully", "success")
                popup.destroy()
                
            except ValueError:
                self.show_message("Please enter valid numbers for quantity and price", "error")
            except Exception as e:
                self.show_message(f"Error updating product: {str(e)}", "error")
        
        # Save button with enhanced styling
        save_button = tk.Button(
            form_frame,
            text="Save Changes",
            command=save_update,
            font=("Segoe UI", 12, "bold")
        )
        SmartStockStyles.apply_button_style(save_button, "primary")
        save_button.pack(fill="x", pady=(20, 0), ipady=12)