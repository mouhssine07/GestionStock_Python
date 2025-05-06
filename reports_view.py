import tkinter as tk
from tkinter import ttk, filedialog
from docx import Document
from fpdf import FPDF
from styles import SmartStockStyles
from datetime import datetime
import os

class ReportsView:
    def __init__(self, parent, products, username):
        self.parent = parent
        self.products = products
        self.username = username
        self.create_reports_ui()
    
    def create_reports_ui(self):
        """Create the reports UI"""
        header_frame = ttk.Frame(self.parent, style="Main.TFrame")
        header_frame.pack(fill="x", padx=30, pady=(30, 20))
        
        title_label = ttk.Label(
            header_frame,
            text="Reports",
            style="Title.TLabel"
        )
        title_label.pack(side="left")
        
        subtitle_label = ttk.Label(
            header_frame,
            text="Generate and view inventory reports",
            font=("Segoe UI", 12),
            background=SmartStockStyles.COLORS["bg_light"],
            foreground=SmartStockStyles.COLORS["text_muted"]
        )
        subtitle_label.pack(side="left", padx=(10, 0), pady=5)
        
        content_frame = ttk.Frame(self.parent, style="Main.TFrame")
        content_frame.pack(fill="both", expand=True, padx=30, pady=(0, 30))
        
        if not self.products:
            no_data_label = ttk.Label(
                content_frame,
                text="No products to report. Add products in the Products section.",
                font=SmartStockStyles.FONTS["body"],
                background=SmartStockStyles.COLORS["bg_light"],
                foreground=SmartStockStyles.COLORS["text_muted"]
            )
            no_data_label.pack(pady=20)
            return
        
        # Report options
        options_frame = ttk.Frame(content_frame, style="Card.TFrame", padding=20)
        options_frame.pack(fill="x", pady=(0, 20))
        
        ttk.Label(
            options_frame,
            text="Select Report Type",
            style="Subtitle.TLabel"
        ).pack(anchor="w", pady=(0, 10))
        
        report_types = ["Inventory Summary", "Low Stock Report"]
        self.report_var = tk.StringVar(value=report_types[0])
        report_menu = ttk.Combobox(
            options_frame,
            textvariable=self.report_var,
            values=report_types,
            state="readonly"
        )
        report_menu.pack(fill="x", pady=(0, 10))
        
        button_frame = ttk.Frame(options_frame, style="Card.TFrame")
        button_frame.pack(fill="x")
        button_frame.columnconfigure(0, weight=1)
        button_frame.columnconfigure(1, weight=1)
        
        word_button = tk.Button(
            button_frame,
            text="Generate Word Report",
            command=lambda: self.generate_report("word")
        )
        SmartStockStyles.apply_button_style(word_button, "primary")
        word_button.grid(row=0, column=0, padx=(0, 5), pady=5, sticky="ew")
        
        pdf_button = tk.Button(
            button_frame,
            text="Generate PDF Report",
            command=lambda: self.generate_report("pdf")
        )
        SmartStockStyles.apply_button_style(pdf_button, "primary")
        pdf_button.grid(row=0, column=1, padx=(5, 0), pady=5, sticky="ew")
        
        # Report preview
        preview_frame = ttk.Frame(content_frame, style="Card.TFrame", padding=20)
        preview_frame.pack(fill="both", expand=True)
        
        ttk.Label(
            preview_frame,
            text="Report Preview",
            style="Subtitle.TLabel"
        ).pack(anchor="w")
        
        self.preview_text = tk.Text(
            preview_frame,
            wrap="word",
            height=10,
            font=SmartStockStyles.FONTS["body"],
            bg=SmartStockStyles.COLORS["bg_card"],
            fg=SmartStockStyles.COLORS["text_dark"],
            relief="flat",
            padx=5,
            pady=5
        )
        self.preview_text.pack(fill="both", expand=True, pady=10)
        self.preview_text.config(state="disabled")
        
        # Update preview on selection change
        report_menu.bind("<<ComboboxSelected>>", self.update_preview)
        self.update_preview(None)
    
    def update_preview(self, event):
        """Update the report preview"""
        report_type = self.report_var.get()
        self.preview_text.config(state="normal")
        self.preview_text.delete("1.0", tk.END)
        
        if report_type == "Inventory Summary":
            total_products = len(self.products)
            total_value = sum(p["quantity"] * p["price"] for p in self.products)
            avg_price = sum(p["price"] for p in self.products) / total_products if total_products else 0
            text = (
                f"Inventory Summary\n"
                f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}\n\n"
                f"Total Products: {total_products}\n"
                f"Total Inventory Value: ${total_value:.2f}\n"
                f"Average Price: ${avg_price:.2f}\n\n"
                f"Top 3 Most Valuable Products:\n"
            )
            sorted_products = sorted(self.products, key=lambda p: p["quantity"] * p["price"], reverse=True)[:3]
            for p in sorted_products:
                text += f"- {p['name']}: {p['quantity']} units, ${p['price']:.2f} each\n"
        else:  # Low Stock Report
            low_stock = [p for p in self.products if p["quantity"] < 5]
            text = (
                f"Low Stock Report\n"
                f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}\n\n"
                f"Products with Low Stock (< 5 units):\n"
            )
            if low_stock:
                for p in low_stock:
                    text += f"- {p['name']}: {p['quantity']} units\n"
            else:
                text += "No products are low on stock."
        
        self.preview_text.insert("1.0", text)
        self.preview_text.config(state="disabled")
    
    def generate_report(self, format_type):
        """Generate a report in Word or PDF format"""
        report_type = self.report_var.get()
        
        if report_type == "Inventory Summary":
            self.generate_inventory_summary(format_type)
        else:  # Low Stock Report
            self.generate_low_stock_report(format_type)
    
    def generate_inventory_summary(self, format_type):
        """Generate an inventory summary report"""
        total_products = len(self.products)
        total_value = sum(p["quantity"] * p["price"] for p in self.products)
        avg_price = sum(p["price"] for p in self.products) / total_products if total_products else 0
        sorted_products = sorted(self.products, key=lambda p: p["quantity"] * p["price"], reverse=True)[:3]
        
        if format_type == "word":
            doc = Document()
            doc.add_heading("Inventory Summary", 0)
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
            doc.add_paragraph(f"User: {self.username}")
            doc.add_paragraph("")
            
            doc.add_paragraph(f"Total Products: {total_products}")
            doc.add_paragraph(f"Total Inventory Value: ${total_value:.2f}")
            doc.add_paragraph(f"Average Price: ${avg_price:.2f}")
            doc.add_paragraph("")
            
            doc.add_paragraph("Top 3 Most Valuable Products:")
            for p in sorted_products:
                doc.add_paragraph(f"- {p['name']}: {p['quantity']} units, ${p['price']:.2f} each")
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Save Inventory Summary As",
                initialfile="inventory_summary.docx"
            )
            
            if file_path:
                doc.save(file_path)
                ttk.Label(self.parent, text=f"Report saved to {os.path.basename(file_path)}", style="Success.TLabel").pack()
        
        else:  # PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(200, 10, txt="Inventory Summary", ln=True, align='C')
            
            pdf.set_font("Arial", '', 10)
            pdf.cell(200, 10, txt=f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}", ln=True)
            pdf.cell(200, 5, txt=f"User: {self.username}", ln=True)
            pdf.ln(5)
            
            pdf.set_font("Arial", '', 12)
            pdf.cell(200, 10, txt=f"Total Products: {total_products}", ln=True)
            pdf.cell(200, 10, txt=f"Total Inventory Value: ${total_value:.2f}", ln=True)
            pdf.cell(200, 10, txt=f"Average Price: ${avg_price:.2f}", ln=True)
            pdf.ln(5)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(200, 10, txt="Top 3 Most Valuable Products:", ln=True)
            pdf.set_font("Arial", '', 12)
            for p in sorted_products:
                pdf.cell(200, 10, txt=f"- {p['name']}: {p['quantity']} units, ${p['price']:.2f} each", ln=True)
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
                title="Save Inventory Summary As",
                initialfile="inventory_summary.pdf"
            )
            
            if file_path:
                pdf.output(file_path)
                ttk.Label(self.parent, text=f"Report saved to {os.path.basename(file_path)}", style="Success.TLabel").pack()
    
    def generate_low_stock_report(self, format_type):
        """Generate a low stock report"""
        low_stock = [p for p in self.products if p["quantity"] < 5]
        
        if format_type == "word":
            doc = Document()
            doc.add_heading("Low Stock Report", 0)
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
            doc.add_paragraph(f"User: {self.username}")
            doc.add_paragraph("")
            
            doc.add_paragraph("Products with Low Stock (< 5 units):")
            if low_stock:
                for p in low_stock:
                    doc.add_paragraph(f"- {p['name']}: {p['quantity']} units")
            else:
                doc.add_paragraph("No products are low on stock.")
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
                title="Save Low Stock Report As",
                initialfile="low_stock_report.docx"
            )
            
            if file_path:
                doc.save(file_path)
                ttk.Label(self.parent, text=f"Report saved to {os.path.basename(file_path)}", style="Success.TLabel").pack()
        
        else:  # PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(200, 10, txt="Low Stock Report", ln=True, align='C')
            
            pdf.set_font("Arial", '', 10)
            pdf.cell(200, 10, txt=f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}", ln=True)
            pdf.cell(200, 5, txt=f"User: {self.username}", ln=True)
            pdf.ln(5)
            
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(200, 10, txt="Products with Low Stock (< 5 units):", ln=True)
            pdf.set_font("Arial", '', 12)
            if low_stock:
                for p in low_stock:
                    pdf.cell(200, 10, txt=f"- {p['name']}: {p['quantity']} units", ln=True)
            else:
                pdf.cell(200, 10, txt="No products are low on stock.", ln=True)
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
                title="Save Low Stock Report As",
                initialfile="low_stock_report.pdf"
            )
            
            if file_path:
                pdf.output(file_path)
                ttk.Label(self.parent, text=f"Report saved to {os.path.basename(file_path)}", style="Success.TLabel").pack()
    
    def update_products(self, products):
        """Update products and refresh UI"""
        self.products = products
        for widget in self.parent.winfo_children():
            widget.destroy()
        self.create_reports_ui()